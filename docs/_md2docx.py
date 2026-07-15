# -*- coding: utf-8 -*-
"""Convierte docs/USER_STORIES.md a .docx con formato (títulos, tablas, negritas)."""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND_GREEN = RGBColor(0x1E, 0x51, 0x26)
BRAND_NIGHT = RGBColor(0x08, 0x22, 0x1A)
CODE_GRAY = RGBColor(0x6B, 0x46, 0x16)

INLINE_RE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+`)')


def add_runs(paragraph, text):
    """Parsea **bold**, *italic*, `code` en runs."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = paragraph.add_run(part[1:-1]); r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9.5); r.font.color.rgb = CODE_GRAY
        else:
            paragraph.add_run(part)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '1E5126')
    pbdr.append(bottom); pPr.append(pbdr)


def main(md_path, out_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'; normal.font.size = Pt(10.5)
    for lvl in range(1, 5):
        try:
            h = doc.styles['Heading %d' % lvl]
            h.font.color.rgb = BRAND_GREEN if lvl > 1 else BRAND_NIGHT
            h.font.name = 'Calibri'
        except KeyError:
            pass

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1; continue

        # Horizontal rule
        if stripped == '---':
            hr(doc); i += 1; continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                p = doc.add_heading('', 0)
                r = p.add_run(text); r.font.color.rgb = BRAND_NIGHT
            else:
                doc.add_heading(text, min(level - 1, 4))
            i += 1; continue

        # Tables: consecutive lines starting with |
        if stripped.startswith('|'):
            block = []
            while i < n and lines[i].strip().startswith('|'):
                block.append(lines[i].strip()); i += 1
            rows = []
            for b in block:
                cells = [c.strip() for c in b.strip('|').split('|')]
                rows.append(cells)
            # drop separator row (---)
            body = [r for r in rows if not all(re.match(r'^:?-{2,}:?$', c) for c in r if c != '') or not r]
            body = [r for r in rows if not (len(r) and all(re.match(r'^:?-{1,}:?$', c.strip()) for c in r))]
            if body:
                ncols = max(len(r) for r in body)
                table = doc.add_table(rows=0, cols=ncols)
                table.style = 'Light Grid Accent 1'
                for ri, r in enumerate(body):
                    cells = table.add_row().cells
                    for ci in range(ncols):
                        txt = r[ci] if ci < len(r) else ''
                        cell = cells[ci]
                        cell.paragraphs[0].text = ''
                        add_runs(cell.paragraphs[0], txt)
                        if ri == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            shade(cell, '1E5126')
            continue

        # Blockquote callout
        if stripped.startswith('>'):
            buf = []
            while i < n and lines[i].strip().startswith('>'):
                buf.append(re.sub(r'^>\s?', '', lines[i].strip())); i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
            add_runs(p, ' '.join(buf))
            for run in p.runs:
                run.italic = True
            # left border bar
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18')
            left.set(qn('w:space'), '8'); left.set(qn('w:color'), 'F58718')
            pbdr.append(left); pPr.append(pbdr)
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if m:
            p = doc.add_paragraph(style='List Number')
            add_runs(p, m.group(2)); i += 1; continue

        # Bullets (one nesting level via 2-space indent)
        m = re.match(r'^(\s*)-\s+(.*)$', line)
        if m:
            indent = len(m.group(1))
            style = 'List Bullet 2' if indent >= 2 else 'List Bullet'
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style='List Bullet')
            add_runs(p, m.group(2)); i += 1; continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(out_path)
    print('Saved:', out_path)


if __name__ == '__main__':
    main(sys.argv[1], sys.argv[2])
