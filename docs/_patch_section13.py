# -*- coding: utf-8 -*-
"""Reemplaza SOLO la sección 13 (matchmaking) en un .docx existente, preservando
todo lo anterior (incluidas las imágenes) y el sectPr final. Usa formato DIRECTO
(no estilos con nombre) para coincidir con el doc aplanado por Google Docs:
viñetas via numPr numId=1, tablas con tblBorders + header verde 1E5126."""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

CODE_GRAY = RGBColor(0x6B, 0x46, 0x16)
INLINE_RE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+`)')


def add_runs(paragraph, text):
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = paragraph.add_run(part[1:-1]); r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9.5); r.font.color.rgb = CODE_GRAY
        else:
            paragraph.add_run(part)


def _set_ind(pPr, left, hanging=None):
    ind = OxmlElement('w:ind'); ind.set(qn('w:left'), str(left))
    if hanging is not None:
        ind.set(qn('w:hanging'), str(hanging))
    pPr.append(ind)


def add_bullet(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '1')
    numPr.append(ilvl); numPr.append(numId); pPr.append(numPr)
    _set_ind(pPr, 360, 360)
    add_runs(p, text)


def add_numbered(doc, num, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    _set_ind(pPr, 360)
    add_runs(p, '%s. %s' % (num, text))


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def add_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), 'CCCCCC')
        borders.append(e)
    tblPr.append(borders)


def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '1E5126')
    pbdr.append(bottom); pPr.append(pbdr)


def callout(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    _set_ind(pPr, 360)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    add_runs(p, text)
    for run in p.runs:
        run.italic = True
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8'); left.set(qn('w:color'), 'F58718')
    pbdr.append(left)
    # pBdr must precede spacing/ind in pPr order -> insert near front
    pPr.insert(0, pbdr)


def render_into(doc, lines):
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]; stripped = line.strip()
        if not stripped:
            i += 1; continue
        if stripped == '---':
            hr(doc); i += 1; continue
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1)); text = m.group(2).strip()
            doc.add_heading(text, 0 if level == 1 else min(level - 1, 4))
            i += 1; continue
        if stripped.startswith('|'):
            block = []
            while i < n and lines[i].strip().startswith('|'):
                block.append(lines[i].strip()); i += 1
            rows = [[c.strip() for c in b.strip('|').split('|')] for b in block]
            body = [r for r in rows if not (len(r) and all(re.match(r'^:?-{1,}:?$', c.strip()) for c in r))]
            if body:
                ncols = max(len(r) for r in body)
                table = doc.add_table(rows=0, cols=ncols)
                table.autofit = True
                add_table_borders(table)
                for ri, r in enumerate(body):
                    cells = table.add_row().cells
                    for ci in range(ncols):
                        txt = r[ci] if ci < len(r) else ''
                        cell = cells[ci]; cell.paragraphs[0].text = ''
                        add_runs(cell.paragraphs[0], txt)
                        if ri == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            shade(cell, '1E5126')
            continue
        if stripped.startswith('>'):
            buf = []
            while i < n and lines[i].strip().startswith('>'):
                buf.append(re.sub(r'^>\s?', '', lines[i].strip())); i += 1
            callout(doc, ' '.join(buf)); continue
        m = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if m:
            add_numbered(doc, m.group(1), m.group(2)); i += 1; continue
        m = re.match(r'^(\s*)-\s+(.*)$', line)
        if m:
            add_bullet(doc, m.group(2)); i += 1; continue
        p = doc.add_paragraph(); add_runs(p, stripped); i += 1


def main(src, md_path, out):
    doc = Document(src)
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))  # keep in place: python-docx inserts before it
    children = list(body)
    cut = None
    for idx, el in enumerate(children):
        if el.tag == qn('w:p'):
            p = Paragraph(el, doc)
            if p.style and p.style.name == 'Heading 1' and p.text.strip().startswith('13.'):
                cut = idx; break
    if cut is None:
        print('ERROR: no encontré el encabezado de la sección 13'); sys.exit(1)
    removed = 0
    for el in children[cut:]:
        if el is sectPr:
            continue  # never remove the section properties
        body.remove(el); removed += 1
    print('Elementos removidos (vieja seccion 13):', removed)
    with open(md_path, encoding='utf-8') as f:
        render_into(doc, f.read().split('\n'))
    doc.save(out)
    print('Guardado:', out)


if __name__ == '__main__':
    main(sys.argv[1], sys.argv[2], sys.argv[3])
